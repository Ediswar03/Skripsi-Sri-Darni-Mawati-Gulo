"""
FINAL CLEAN REBUILD:
Remove ALL injected paragraphs (DAFTAR ISI entries + DAFTAR TABEL + DAFTAR GAMBAR)
Then insert them in the correct order using XML append (not insert_before which reverses order).
"""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import copy

DOC_PATH = r"g:\Skripsi Sri Darni\File Skripsi\Sri Darni.docx"
doc = Document(DOC_PATH)

# ============================================================
# Step 1: Find and delete ALL injected front matter (117 to 193)
# ============================================================
# Injected content starts at a line with '. . . . . i' (ABSTRAK entry)
# and ends at the bare 'DAFTAR ISI' heading (original, at para 175 after prev edit)

inject_start = None
daftar_isi_para_idx = None

for i, p in enumerate(doc.paragraphs):
    if inject_start is None and ('. . . . . i' in p.text or '. . . . . ii' in p.text or '. . . . . iii' in p.text):
        inject_start = i
    if inject_start is not None and p.text.strip() == 'DAFTAR ISI':
        daftar_isi_para_idx = i
        break

print(f"Injected content starts at: {inject_start}")
print(f"DAFTAR ISI heading at: {daftar_isi_para_idx}")

if inject_start and daftar_isi_para_idx:
    # Delete all injected paragraphs before DAFTAR ISI heading
    to_delete = list(range(inject_start, daftar_isi_para_idx))
    print(f"Deleting {len(to_delete)} injected paragraphs...")
    for idx in reversed(to_delete):
        p = doc.paragraphs[idx]
        p._element.getparent().remove(p._element)
    print("Done.")

# Also delete 'Disiplin kerja' point from identifikasi masalah if still present
for p in doc.paragraphs:
    if 'Disiplin kerja berpengaruh terhadap kinerja perangkat desa' in p.text:
        p._element.getparent().remove(p._element)
        print("Removed: Disiplin kerja point")
        break

# ============================================================
# Step 2: Find the DAFTAR ISI heading paragraph (now at a lower index)
# ============================================================
daftar_isi_para = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == 'DAFTAR ISI':
        daftar_isi_para = p
        print(f"DAFTAR ISI heading now at paragraph {i}")
        break

# ============================================================
# Step 3: Build XML elements for all front matter and append
#         in correct order AFTER the DAFTAR ISI heading
# ============================================================

def make_para_xml(text, font_size=12, bold=False, italic=False, center=False, dots_target=None, page=None):
    """Create a w:p XML element with given text and formatting."""
    p_el = OxmlElement('w:p')
    
    # Paragraph properties
    pPr = OxmlElement('w:pPr')
    if center:
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'center')
        pPr.append(jc)
    # Line spacing 1.5
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:line'), '360')   # 360 = 1.5 lines (240 = single)
    spacing.set(qn('w:lineRule'), 'auto')
    spacing.set(qn('w:after'), '40')
    pPr.append(spacing)
    p_el.append(pPr)
    
    # Build display text
    if page is not None:
        dots_count = max(3, 70 - len(text))
        dots = ' .' * (dots_count // 2)
        display = f"{text}{dots}  {page}"
    else:
        display = text
    
    # Run
    r_el = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    # Font
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rPr.append(rFonts)
    # Size (in half-points, so 12pt = 24)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(font_size * 2))
    rPr.append(sz)
    if bold:
        b_el = OxmlElement('w:b')
        rPr.append(b_el)
    if italic:
        i_el = OxmlElement('w:i')
        rPr.append(i_el)
    r_el.append(rPr)
    
    t_el = OxmlElement('w:t')
    t_el.text = display
    t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r_el.append(t_el)
    p_el.append(r_el)
    return p_el

def make_page_break_xml():
    """Create a paragraph with a page break."""
    p_el = OxmlElement('w:p')
    r_el = OxmlElement('w:r')
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r_el.append(br)
    p_el.append(r_el)
    return p_el

def make_section_header_xml(title):
    """Create a bold, centered section header paragraph."""
    return make_para_xml(title, bold=True, center=True)

# ============================================================
# DAFTAR ISI entries (in correct order: ABSTRAK first, then rest)
# ============================================================
daftar_isi_entries = [
    ("ABSTRAK", "i", True),
    ("ABSTRACT", "ii", True),
    ("KATA PENGANTAR", "iii", True),
    ("DAFTAR ISI", "iv", True),
    ("DAFTAR TABEL", "vi", True),
    ("DAFTAR GAMBAR", "vii", True),
    ("BAB I PENDAHULUAN", "1", True),
    ("    A. Latar Belakang Masalah", "1", False),
    ("    B. Identifikasi Masalah", "5", False),
    ("    C. Pembatasan Masalah", "6", False),
    ("    D. Rumusan Masalah", "6", False),
    ("    E. Tujuan Penelitian", "7", False),
    ("    F. Manfaat Penelitian", "7", False),
    ("    G. Sistematika Penulisan Skripsi", "8", False),
    ("BAB II LANDASAN TEORITIS DAN PENGEMBANGAN HIPOTESIS", "10", True),
    ("    A. Landasan Teori", "10", False),
    ("        1. Kinerja Perangkat Desa", "10", False),
    ("        2. Gaya Kepemimpinan", "13", False),
    ("        3. Kompensasi", "16", False),
    ("        4. Motivasi Kerja", "19", False),
    ("    B. Hasil Penelitian Terdahulu", "23", False),
    ("    C. Kerangka Berpikir", "29", False),
    ("    D. Pengembangan Hipotesis", "30", False),
    ("BAB III METODE PENELITIAN", "31", True),
    ("    A. Tempat dan Waktu Penelitian", "31", False),
    ("    B. Desain Penelitian", "31", False),
    ("    C. Definisi Operasional dan Pengukuran Variabel Penelitian", "32", False),
    ("        1. Variabel Bebas (Variabel Independen)", "32", False),
    ("        2. Variabel Terikat (Variabel Dependen)", "32", False),
    ("        3. Variabel Mediasi", "33", False),
    ("        4. Operasionalisasi Variabel Penelitian", "33", False),
    ("    D. Populasi dan Sampel Penelitian", "34", False),
    ("        1. Populasi", "34", False),
    ("        2. Sampel", "35", False),
    ("    E. Metode Pengumpulan Data", "36", False),
    ("        1. Teknik Pengumpulan Data", "36", False),
    ("        2. Sumber Data", "36", False),
    ("        3. Teknik Pengembangan Instrumen", "37", False),
    ("    F. Teknik Analisis Data", "38", False),
    ("DAFTAR PUSTAKA", "44", True),
]

daftar_tabel_entries = [
    "Tabel 2.1 Faktor-Faktor yang Mempengaruhi Kompensasi.............18",
    "Tabel 2.2 Hasil Penelitian Terdahulu.............................24",
    "Tabel 3.1 Jadwal Penelitian dan Penulisan Skripsi................31",
    "Tabel 3.2 Operasional Variabel Penelitian........................33",
    "Tabel 3.3 Skala Likert...........................................37",
    "Tabel 3.4 Rule of Thumb Evaluasi Model Pengukuran................40",
    "Tabel 3.5 Rule of Thumb Evaluasi Model Struktural................43",
]

daftar_gambar_entries = [
    "Gambar 2.1 Kerangka Berpikir....................................29",
    "Gambar 3.1 Outer Model Variabel Dependen (Y).....................38",
    "Gambar 3.2 Outer Model Variabel Independen (X1)..................38",
    "Gambar 3.3 Outer Model Variabel Independen (X2)..................39",
    "Gambar 3.4 Outer Model Variabel Mediasi (Z)......................39",
    "Gambar 3.5 Inner Model Variabel Dependen, Independen, Mediasi...39",
    "Gambar 3.6 Diagram Jalur (Path Diagram)..........................40",
]

# Get the parent body element and position after DAFTAR ISI heading
body = daftar_isi_para._element.getparent()
di_idx = list(body).index(daftar_isi_para._element)

# We'll insert all elements AFTER the DAFTAR ISI heading
insert_after = di_idx  # position to insert after

def append_after(el, position_ref):
    """Append el after the element at position_ref in body."""
    body.insert(list(body).index(position_ref) + 1, el)
    return el

# Build all XML and insert in correct order
# We keep track of the "last inserted" to always append after it
current_ref = daftar_isi_para._element

for entry_text, page, is_bold in daftar_isi_entries:
    p_el = make_para_xml(entry_text, bold=is_bold, page=page)
    body.insert(list(body).index(current_ref) + 1, p_el)
    current_ref = p_el

print("DAFTAR ISI entries inserted.")

# Page break then DAFTAR TABEL
pb = make_page_break_xml()
body.insert(list(body).index(current_ref) + 1, pb)
current_ref = pb

h_dt = make_section_header_xml("DAFTAR TABEL")
body.insert(list(body).index(current_ref) + 1, h_dt)
current_ref = h_dt

for entry_line in daftar_tabel_entries:
    p_el = make_para_xml(entry_line)
    body.insert(list(body).index(current_ref) + 1, p_el)
    current_ref = p_el

print("DAFTAR TABEL entries inserted.")

# Page break then DAFTAR GAMBAR
pb2 = make_page_break_xml()
body.insert(list(body).index(current_ref) + 1, pb2)
current_ref = pb2

h_dg = make_section_header_xml("DAFTAR GAMBAR")
body.insert(list(body).index(current_ref) + 1, h_dg)
current_ref = h_dg

for entry_line in daftar_gambar_entries:
    p_el = make_para_xml(entry_line)
    body.insert(list(body).index(current_ref) + 1, p_el)
    current_ref = p_el

print("DAFTAR GAMBAR entries inserted.")

doc.save(DOC_PATH)
print("Document saved successfully!")

# Quick verify
doc2 = Document(DOC_PATH)
print(f"\nFinal paragraph count: {len(doc2.paragraphs)}")
for i, p in enumerate(doc2.paragraphs[112:130]):
    txt = p.text.strip()
    if txt:
        print(f"[{i+112}] {txt[:80]}")
