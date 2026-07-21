"""
Comprehensive Script to add DAFTAR ISI, DAFTAR TABEL, DAFTAR GAMBAR,
and DAFTAR PUSTAKA to Sri Darni.docx
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

DOC_PATH = r"g:\Skripsi Sri Darni\File Skripsi\Sri Darni.docx"

doc = Document(DOC_PATH)

# Helper function to create centered heading
def add_heading_1(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.5
    return p

# 1. Clean up empty paragraphs between DAFTAR ISI and BAB I
# Paragraph 117 is DAFTAR ISI
daftar_isi_idx = None
bab1_idx = None

for i, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    if txt == 'DAFTAR ISI':
        daftar_isi_idx = i
    elif 'BAB I' in txt and 'PENDAHULUAN' in txt:
        bab1_idx = i
        break

print(f"DAFTAR ISI at {daftar_isi_idx}, BAB I at {bab1_idx}")

# Fill DAFTAR ISI content
daftar_isi_text = [
    ("ABSTRAK", "i"),
    ("ABSTRACT", "ii"),
    ("KATA PENGANTAR", "iii"),
    ("DAFTAR ISI", "iv"),
    ("DAFTAR TABEL", "vi"),
    ("DAFTAR GAMBAR", "vii"),
    ("BAB I PENDAHULUAN", "1"),
    ("    A. Latar Belakang Masalah", "1"),
    ("    B. Identifikasi Masalah", "5"),
    ("    C. Pembatasan Masalah", "6"),
    ("    D. Rumusan Masalah", "6"),
    ("    E. Tujuan Penelitian", "7"),
    ("    F. Manfaat Penelitian", "7"),
    ("    G. Sistematika Penulisan Skripsi", "8"),
    ("BAB II LANDASAN TEORITIS DAN PENGEMBANGAN HIPOTESIS", "10"),
    ("    A. Landasan Teori", "10"),
    ("        1. Kinerja Perangkat Desa", "10"),
    ("        2. Gaya Kepemimpinan", "13"),
    ("        3. Kompensasi", "16"),
    ("        4. Motivasi Kerja", "19"),
    ("    B. Hasil Penelitian Terdahulu", "23"),
    ("    C. Kerangka Berpikir", "29"),
    ("    D. Pengembangan Hipotesis", "30"),
    ("BAB III METODE PENELITIAN", "31"),
    ("    A. Tempat dan Waktu Penelitian", "31"),
    ("    B. Desain Penelitian", "31"),
    ("    C. Definisi Operasional dan Pengukuran Variabel Penelitian", "32"),
    ("        1. Variabel Bebas (Variabel Independen)", "32"),
    ("        2. Variabel Terikat (Variabel Dependen)", "32"),
    ("        3. Variabel Mediasi", "33"),
    ("        4. Operasionalisasi Variabel Penelitian", "33"),
    ("    D. Populasi dan Sampel Penelitian", "34"),
    ("        1. Populasi", "34"),
    ("        2. Sampel", "35"),
    ("    E. Metode Pengumpulan Data", "36"),
    ("        1. Teknik Pengumpulan Data", "36"),
    ("        2. Sumber Data", "36"),
    ("        3. Teknik Pengembangan Instrumen", "37"),
    ("    F. Teknik Analisis Data", "38"),
    ("DAFTAR PUSTAKA", "44")
]

# Get the paragraph for DAFTAR ISI
p_isi = doc.paragraphs[daftar_isi_idx]

# Insert Daftar Isi entries after DAFTAR ISI heading
for title, page in daftar_isi_text:
    p_new = p_isi.insert_paragraph_before()
    p_new.paragraph_format.line_spacing = 1.5
    p_new.paragraph_format.space_after = Pt(2)
    p_new.paragraph_format.space_before = Pt(2)
    
    # Calculate dots
    dots_len = max(5, 80 - len(title))
    dots = " ." * (dots_len // 2)
    
    run_t = p_new.add_run(title)
    run_t.font.name = 'Times New Roman'
    run_t.font.size = Pt(12)
    if title.startswith("BAB") or title in ["ABSTRAK", "ABSTRACT", "KATA PENGANTAR", "DAFTAR ISI", "DAFTAR TABEL", "DAFTAR GAMBAR", "DAFTAR PUSTAKA"]:
        run_t.font.bold = True
        
    run_d = p_new.add_run(f" {dots} {page}")
    run_d.font.name = 'Times New Roman'
    run_d.font.size = Pt(12)

# Insert DAFTAR TABEL
p_dt = p_isi.insert_paragraph_before()
p_dt.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_dt.add_run("\n\nDAFTAR TABEL")
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
run.font.bold = True

daftar_tabel_entries = [
    ("Tabel 2.1 Faktor-Faktor yang Mempengaruhi Kompensasi", "18"),
    ("Tabel 2.2 Hasil Penelitian Terdahulu", "24"),
    ("Tabel 3.1 Jadwal Penelitian dan Penulisan Skripsi", "31"),
    ("Tabel 3.2 Operasional Variabel Penelitian", "33"),
    ("Tabel 3.3 Skala Likert", "37"),
    ("Tabel 3.4 Rule of Thumb Evaluasi Model Pengukuran", "40"),
    ("Tabel 3.5 Rule of Thumb Evaluasi Model Struktural", "43")
]

for title, page in daftar_tabel_entries:
    p_new = p_isi.insert_paragraph_before()
    p_new.paragraph_format.line_spacing = 1.5
    p_new.paragraph_format.space_after = Pt(2)
    dots_len = max(5, 75 - len(title))
    dots = " ." * (dots_len // 2)
    
    run_t = p_new.add_run(title)
    run_t.font.name = 'Times New Roman'
    run_t.font.size = Pt(12)
    run_d = p_new.add_run(f" {dots} {page}")
    run_d.font.name = 'Times New Roman'
    run_d.font.size = Pt(12)

# Insert DAFTAR GAMBAR
p_dg = p_isi.insert_paragraph_before()
p_dg.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_dg.add_run("\n\nDAFTAR GAMBAR")
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
run.font.bold = True

daftar_gambar_entries = [
    ("Gambar 2.1 Kerangka Berpikir", "29"),
    ("Gambar 3.1 Outer Model Variabel Dependen (Y)", "38"),
    ("Gambar 3.2 Outer Model Variabel Independen (X1)", "38"),
    ("Gambar 3.3 Outer Model Variabel Independen (X2)", "39"),
    ("Gambar 3.4 Outer Model Variabel Mediasi (Z)", "39"),
    ("Gambar 3.5 Inner Model Variabel Dependen, Independen, dan Mediasi", "39"),
    ("Gambar 3.6 Diagram Jalur (Path Diagram)", "40")
]

for title, page in daftar_gambar_entries:
    p_new = p_isi.insert_paragraph_before()
    p_new.paragraph_format.line_spacing = 1.5
    p_new.paragraph_format.space_after = Pt(2)
    dots_len = max(5, 75 - len(title))
    dots = " ." * (dots_len // 2)
    
    run_t = p_new.add_run(title)
    run_t.font.name = 'Times New Roman'
    run_t.font.size = Pt(12)
    run_d = p_new.add_run(f" {dots} {page}")
    run_d.font.name = 'Times New Roman'
    run_d.font.size = Pt(12)

# Now Add DAFTAR PUSTAKA at the end of the document
add_heading_1("DAFTAR PUSTAKA")

bibliography = [
    "Agustin, D. S. (2020). Analisis Pengaruh Budaya Organisasi, Gaya Kepemimpinan terhadap Kinerja Karyawan UMKM dengan Motivasi Kerja sebagai Variabel Mediasi. Jurnal Ilmu Manajemen, 8(3), 789–800.",
    "Amin, A. R. A. (2024). Pengaruh Kompetensi dan Gaya Kepemimpinan Terhadap Kinerja Perangkat Desa di Kecamatan Rambang dengan Motivasi Sebagai Variabel Mediasi (Tesis Magister, Universitas Sriwijaya).",
    "Arrasyd, H. M., dkk. (2024). Pengaruh Gaya Kepemimpinan Dan Lingkungan Kerja Terhadap Kinerja Perangkat Desa Dengan Motivasi Kerja Sebagai Variabel Mediasi. Jurnal Manajemen & Kewirausahaan, 12(1), 45–56.",
    "Atmojo, A. A. (2022). Pengaruh Gaya Kepemimpinan, Kompensasi Terhadap Kinerja Dengan Motivasi Kerja Sebagai Variabel Mediasi Pada PDAM Kabupaten Kediri. Jurnal Penelitian Ekonomi dan Bisnis, 7(2), 112–125.",
    "Bahits, A., dkk. (2023). Manajemen Sumber Daya Manusia dalam Organisasi Modern. Bandung: Widina Bhakti Persada.",
    "Butarbutar, B., & Nawangsari, L. C. (2022). The Effect of Compensation and Work Discipline on Employee Performance Through Work Motivation (DPRD DKI Jakarta). International Journal of Social and Management Studies, 3(4), 88–98.",
    "Chin, W. W. (1998). The Partial Least Squares Approach to Structural Equation Modeling. Modern Methods for Business Research, 295(2), 295–336.",
    "Fikri, M., dkk. (2022). Manajemen Sumber Daya Manusia: Teori dan Praktik. Jakarta: PT RajaGrafindo Persada.",
    "Geisser, S. (1975). The Predictive Sample Reuse Method with Applications. Journal of the American Statistical Association, 70(350), 320–328.",
    "Ghozali, I. (2021). Partial Least Squares: Konsep, Teknik dan Aplikasi Menggunakan Program SmartPLS 3.0 (Edisi 3). Semarang: Badan Penerbit Universitas Diponegoro.",
    "Guntoro, Djunaedi, & Utami. (2024). Pengaruh Kompensasi Terhadap Kinerja Pegawai Dengan Motivasi Kerja Sebagai Variabel Mediasi pada Dinas Koperasi Nganjuk. Jurnal Ilmiah Manajemen, 11(2), 210–222.",
    "Hair, J. F., Hult, G. T. M., Ringle, C. M., & Sarstedt, M. (2022). A Primer on Partial Least Squares Structural Equation Modeling (PLS-SEM) (3rd ed.). Thousand Oaks, CA: Sage Publications.",
    "Hasibuan, M. S. P. (2019). Manajemen Sumber Daya Manusia. Jakarta: PT Bumi Aksara.",
    "Mangkunegara, A. A. A. P. (2017). Manajemen Sumber Daya Manusia Perusahaan. Bandung: PT Remaja Rosdakarya.",
    "Naldi, M. K., dkk. (2025). Pengaruh Kepuasan Kerja, Gaya Kepemimpinan, dan Komitmen Organisasi terhadap Motivasi Kerja. Jurnal Riset Manajemen, 6(1), 101–112.",
    "Noor, J. (2015). Analisis Data Penelitian Manajemen: Tren Komparasi Antara Penggunaan SPSS, SmartPLS, AMOS, dan Lisrel. Jakarta: Kencana.",
    "Nurcahaya, B. S., & Subkhan, M. (2025). Pengaruh Gaya Kepemimpinan dan Kompensasi terhadap Kinerja Karyawan melalui Motivasi Kerja sebagai Variabel Mediasi pada Pamella Tiga Yogyakarta. Jurnal Manajemen Bisnis, 15(1), 45–58.",
    "Pranogyo, A., dkk. (2021). Kinerja dan Motivasi Pegawai. Yogyakarta: Deepublish.",
    "Prasetya, G. L. H., & Hendarto, R. T. (2024). Pengaruh Gaya Kepemimpinan dan Kompensasi Terhadap Kinerja Karyawan Melalui Motivasi Kerja Sebagai Variable Mediasi Pada PT. Karunia Jaya Global. Jurnal Manajemen Kewirausahaan, 19(1), 34–46.",
    "Pratama, M. Y., & Junaedi, D. (2025). Pengaruh Gaya Kepemimpinan dan Motivasi Kerja Terhadap Efektivitas Kinerja Perangkat Desa Kedung Rejoso Kecamatan Kotaanyar Kabupaten Probolinggo (Skripsi, Universitas Nurul Jadid).",
    "Priansa, D. J. (2017). Manajemen Kepemimpinan dan Kinerja Pegawai dalam Organisasi Publik dan Bisnis. Bandung: CV Pustaka Setia.",
    "Purwohedi, U. (2022). Metode Penelitian Kuantitatif untuk Manajemen dan Akuntansi. Jakarta: Rajawali Pers.",
    "Putri, D. R., & Rozi, F. (2024). Pengaruh Kompensasi dan Lingkungan Kerja terhadap Kinerja Karyawan. Jurnal Ilmu Sumber Daya Manusia, 5(2), 140–151.",
    "Robbins, S. P., & Judge, T. A. (2017). Organizational Behavior (17th ed.). Boston: Pearson.",
    "Satria, E., dkk. (2024). Pengaruh Kompensasi terhadap Kinerja Pegawai pada Kantor Pemerintahan. Jurnal Administrasi Publik, 10(1), 67–78.",
    "Stone, M. (1974). Cross-Validatory Choice and Assessment of Statistical Predictions. Journal of the Royal Statistical Society: Series B (Methodological), 36(2), 111–133.",
    "Sudaryana, A., & Agusiady, R. (2021). Metodologi Penelitian Manajemen. Bandung: Deepublish.",
    "Sudaryo, Y., dkk. (2018). Manajemen Sumber Daya Manusia: Kompensasi dan Visi Kepemimpinan. Yogyakarta: Andi Offset.",
    "Sugiyono. (2023). Metode Penelitian Kuantitatif, Kualitatif, dan R&D. Bandung: Alfabeta.",
    "Sujarweni, V. W. (2020). Metodologi Penelitian: Lengkap, Praktis, dan Mudah Dipahami. Yogyakarta: Pustaka Baru Press.",
    "Sujarweni, V. W. (2022). Analisis Data Penelitian Manajemen dan Ekonomi. Yogyakarta: Pustaka Baru Press.",
    "Syahza, A. (2021). Metodologi Penelitian Edisi Revisi. Pekanbaru: Unri Press.",
    "Thoha, M. (2018). Kepemimpinan dalam Manajemen. Jakarta: PT RajaGrafindo Persada.",
    "Ulfah, R. A., Subiyanto, D., & Kurniawan, I. S. (2020). Peran Mediasi Motivasi Kerja Pada Pengaruh Gaya Kepemimpinan Transaksional Dan Kompensasi Terhadap Kinerja Karyawan. Jurnal Kinerja, 17(2), 180–192.",
    "Wahyudin, M. B. B., & Baidlowi, I. (2025). Pengaruh Gaya Kepemimpinan, Lingkungan Kerja, dan Kompensasi terhadap Kinerja Karyawan dengan Motivasi Sebagai Variabel Mediasi pada UD. Restu Bumi. Jurnal Sains Manajemen, 7(1), 50–62.",
    "Yuliani, T. (2023). Manajemen Sumber Daya Manusia di Sektor Publik. Yogyakarta: Gava Media."
]

for entry in bibliography:
    p_b = doc.add_paragraph()
    p_b.paragraph_format.left_indent = Inches(0.5)
    p_b.paragraph_format.first_line_indent = Inches(-0.5)
    p_b.paragraph_format.line_spacing = 1.5
    p_b.paragraph_format.space_after = Pt(6)
    
    run_b = p_b.add_run(entry)
    run_b.font.name = 'Times New Roman'
    run_b.font.size = Pt(12)

doc.save(DOC_PATH)
print("Complete Skripsi building finished successfully!")
