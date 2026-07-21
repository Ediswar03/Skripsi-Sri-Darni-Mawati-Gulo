"""
Script to add 3 new rows to Tabel 2.2 (Penelitian Terdahulu) in Sri Darni.docx
Uses python-docx low-level XML manipulation to copy last row as template
"""
import copy
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

DOC_PATH = r"g:\Skripsi Sri Darni\File Skripsi\Sri Darni.docx"

doc = Document(DOC_PATH)

# Table index 1 = Table 2 (Tabel 2.2 Penelitian Terdahulu)
tbl = doc.tables[1]
print(f"Table has {len(tbl.rows)} rows and {len(tbl.columns)} cols")

# New rows data to add
new_rows_data = [
    {
        "col1": "Naldi, M. K., dkk. (2025)",
        "col2": "Pengaruh Kepuasan Kerja, Gaya Kepemimpinan, dan Komitmen Organisasi terhadap Motivasi Kerja",
        "col3": "Motivasi Kerja (Y)",
        "col4": "Gaya Kepemimpinan (X1)",
        "col5": "-",
        "col6": "Gaya kepemimpinan berpengaruh positif dan signifikan terhadap motivasi kerja karyawan.",
        "col7": "Perbedaan terletak pada variabel terikat yang hanya meneliti motivasi kerja, bukan kinerja. Penelitian ini tidak menggunakan variabel mediasi."
    },
    {
        "col1": "Putri, D. R., & Rozi, F. (2024)",
        "col2": "Pengaruh Kompensasi dan Lingkungan Kerja terhadap Kinerja Karyawan",
        "col3": "Kinerja (Y)",
        "col4": "Kompensasi (X1)",
        "col5": "-",
        "col6": "Kompensasi berpengaruh positif dan signifikan terhadap kinerja karyawan.",
        "col7": "Perbedaan terletak pada variabel bebas tambahan (lingkungan kerja) dan penelitian ini tidak menggunakan variabel mediasi."
    },
    {
        "col1": "Satria, E., dkk. (2024)",
        "col2": "Pengaruh Kompensasi terhadap Kinerja Pegawai pada Kantor Pemerintahan",
        "col3": "Kinerja (Y)",
        "col4": "Kompensasi (X1)",
        "col5": "-",
        "col6": "Kompensasi berpengaruh positif dan signifikan terhadap kinerja pegawai pada instansi pemerintah.",
        "col7": "Perbedaan terletak pada obyek penelitian instansi pemerintah umum, bukan perangkat desa, dan tidak menggunakan variabel mediasi."
    }
]

def copy_row_and_clear(source_row):
    """Deep copy a row and clear all text content from the copy"""
    new_tr = copy.deepcopy(source_row._tr)
    # Clear text from all cells in the copy
    for tc in new_tr.findall('.//' + qn('w:tc')):
        for p in tc.findall('.//' + qn('w:p')):
            for r in p.findall('.//' + qn('w:r')):
                for t in r.findall(qn('w:t')):
                    t.text = ''
    return new_tr

def set_cell_text(cell, text):
    """Set the text of a cell, preserving paragraph formatting"""
    # Get first paragraph and clear its runs
    p = cell.paragraphs[0]
    for run in p.runs:
        run.text = ''
    # Set text in first run or add new run
    if p.runs:
        p.runs[0].text = text
    else:
        p.add_run(text)

# Get last data row (row 12, index 11) as template for new rows
last_data_row = tbl.rows[11]
print(f"Last row col1: {last_data_row.cells[0].text[:40]}")

# Get the table's XML element
tbl_element = tbl._tbl

# Add 3 new rows by copying and appending
for i, row_data in enumerate(new_rows_data):
    # Deep copy the last data row structure (preserves formatting, borders, etc.)
    new_tr = copy_row_and_clear(last_data_row)
    
    # Append the new row to the table
    tbl_element.append(new_tr)
    
    print(f"Row {13+i} XML appended.")

# Re-access the table to get new rows
tbl = doc.tables[1]
new_row_count = len(tbl.rows)
print(f"New row count: {new_row_count}")

# Fill the 3 new rows (they are at indices 12, 13, 14)
for i, row_data in enumerate(new_rows_data):
    row_idx = 12 + i
    if row_idx < len(tbl.rows):
        row = tbl.rows[row_idx]
        cells = row.cells
        
        # Set text for each column
        values = [row_data["col1"], row_data["col2"], row_data["col3"],
                  row_data["col4"], row_data["col5"], row_data["col6"], row_data["col7"]]
        
        for col_idx, val in enumerate(values):
            if col_idx < len(cells):
                try:
                    set_cell_text(cells[col_idx], val)
                except Exception as e:
                    print(f"  Warning on cell [{row_idx},{col_idx}]: {e}")
        
        print(f"Row {row_idx+1} filled: {row_data['col1'][:30]}")

# Save the document
doc.save(DOC_PATH)
print("Document saved successfully!")
print(f"Final table has {len(doc.tables[1].rows)} rows")
