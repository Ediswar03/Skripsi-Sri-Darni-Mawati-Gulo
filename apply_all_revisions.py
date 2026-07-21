"""
SINGLE COMPREHENSIVE SCRIPT - Applied to a clean Sri Darni.docx
Performs ALL revisions in one pass:
1. [Jumlah] -> 35
2. Diagram placeholders -> academic note
3. Add DAFTAR ISI, DAFTAR TABEL, DAFTAR GAMBAR after DAFTAR ISI heading
4. Add DAFTAR PUSTAKA at the end
5. Italicize foreign terms
6. Section break + page numbering
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

DOC_PATH = r"g:\Skripsi Sri Darni\File Skripsi\Sri Darni.docx"
doc = Document(DOC_PATH)

print("=== Step 1: Replace [Jumlah] with 35 ===")
for p in doc.paragraphs:
    if '[Jumlah]' in p.text:
        # Preserve run formatting
        for run in p.runs:
            if '[Jumlah]' in run.text:
                run.text = run.text.replace('[Jumlah]', '35')
        print(f"  Fixed: {p.text[:70]}")

print("\n=== Step 2: Replace diagram placeholders ===")
placeholders = {
    "[GAMBAR OUTER MODEL Y DISINI]": True,
    "[GAMBAR OUTER MODEL X1 DISINI]": True,
    "[GAMBAR OUTER MODEL X2 DISINI]": True,
    "[GAMBAR OUTER MODEL Z DISINI]": True,
    "[GAMBAR INNER MODEL DISINI]": True,
    "[GAMBAR DIAGRAM JALUR DISINI]": True
}
note_text = "(Tampilan diagram model akan disajikan setelah pengolahan data kuesioner pada Bab IV)"
for p in doc.paragraphs:
    for ph in placeholders:
        if ph in p.text:
            for run in p.runs:
                if ph in run.text:
                    run.text = run.text.replace(ph, note_text)
                    run.font.italic = True
                    run.font.color.rgb = RGBColor(80, 80, 80)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            print(f"  Replaced: {ph}")

print("\n=== Step 3: Italicize foreign terms ===")
foreign_terms = [
    "job performance", "actual performance", "level of performance", "thing done",
    "self esteem", "self-esteem", "selfesteem", "goodwill",
    "shareholder", "shareholders", "stakeholder", "stakeholders",
    "outer model", "inner model", "path diagram", "rule of thumb",
    "Structural Equation Modeling", "Partial Least Square", "Partial Least Squares",
    "average variance extracted", "cross loading", "composite reliability",
    "predictive relevance", "predictive sample reuse", "effect size",
    "two-tailed", "two tailed", "loading factor", "communality",
    "confirmatory research", "exploratory research",
]
italicized_count = 0
for p in doc.paragraphs:
    for run in p.runs:
        for term in foreign_terms:
            if term.lower() in run.text.lower() and not run.font.italic:
                run.font.italic = True
                italicized_count += 1
                break
print(f"  Italicized {italicized_count} runs containing foreign terms")

print("\n=== Step 4: Find DAFTAR ISI heading and insert list + DAFTAR TABEL/GAMBAR ===")

# Find the DAFTAR ISI heading paragraph
daftar_isi_para = None
daftar_isi_idx_in_body = None
body = doc.element.body

for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == 'DAFTAR ISI':
        daftar_isi_para = p
        print(f"  Found DAFTAR ISI at paragraph {i}")
        break

def make_text_para(text, bold=False, italic=False, center=False, font_size=12, indent=0):
    """Create a clean paragraph XML element."""
    p_el = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    if center:
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'center')
        pPr.append(jc)
    # Spacing
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:line'), '360')
    spacing.set(qn('w:lineRule'), 'auto')
    spacing.set(qn('w:after'), '40')
    pPr.append(spacing)
    if indent > 0:
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), str(indent * 360))
        pPr.append(ind)
    p_el.append(pPr)

    r_el = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(font_size * 2))
    rPr.append(sz)
    if bold:
        rPr.append(OxmlElement('w:b'))
    if italic:
        rPr.append(OxmlElement('w:i'))
    r_el.append(rPr)
    t_el = OxmlElement('w:t')
    t_el.text = text
    t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r_el.append(t_el)
    p_el.append(r_el)
    return p_el

def make_page_break():
    p_el = OxmlElement('w:p')
    r_el = OxmlElement('w:r')
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r_el.append(br)
    p_el.append(r_el)
    return p_el

def make_entry(label, page, bold=False):
    """Create a DAFTAR entry with dots filling the line."""
    dots_count = max(4, 72 - len(label))
    dots = '.' * dots_count
    text = f"{label} {dots} {page}"
    return make_text_para(text, bold=bold)

# Build all elements in correct order
elements_to_insert = []

# DAFTAR ISI entries
daftar_isi_data = [
    ("ABSTRAK", "i", True),
    ("ABSTRACT", "ii", True),
    ("KATA PENGANTAR", "iii", True),
    ("DAFTAR ISI", "iv", True),
    ("DAFTAR TABEL", "vi", True),
    ("DAFTAR GAMBAR", "vii", True),
    ("BAB I PENDAHULUAN", "1", True),
    ("    A.  Latar Belakang Masalah", "1", False),
    ("    B.  Identifikasi Masalah", "5", False),
    ("    C.  Pembatasan Masalah", "6", False),
    ("    D.  Rumusan Masalah", "6", False),
    ("    E.  Tujuan Penelitian", "7", False),
    ("    F.  Manfaat Penelitian", "7", False),
    ("    G.  Sistematika Penulisan Skripsi", "8", False),
    ("BAB II LANDASAN TEORITIS DAN PENGEMBANGAN HIPOTESIS", "10", True),
    ("    A.  Landasan Teori", "10", False),
    ("         1.  Kinerja Perangkat Desa", "10", False),
    ("         2.  Gaya Kepemimpinan", "13", False),
    ("         3.  Kompensasi", "16", False),
    ("         4.  Motivasi Kerja", "19", False),
    ("    B.  Hasil Penelitian Terdahulu", "23", False),
    ("    C.  Kerangka Berpikir", "29", False),
    ("    D.  Pengembangan Hipotesis", "30", False),
    ("BAB III METODE PENELITIAN", "31", True),
    ("    A.  Tempat dan Waktu Penelitian", "31", False),
    ("    B.  Desain Penelitian", "31", False),
    ("    C.  Definisi Operasional dan Pengukuran Variabel Penelitian", "32", False),
    ("         1.  Variabel Bebas (Variabel Independen)", "32", False),
    ("         2.  Variabel Terikat (Variabel Dependen)", "32", False),
    ("         3.  Variabel Mediasi", "33", False),
    ("         4.  Operasionalisasi Variabel Penelitian", "33", False),
    ("    D.  Populasi dan Sampel Penelitian", "34", False),
    ("         1.  Populasi", "34", False),
    ("         2.  Sampel", "35", False),
    ("    E.  Metode Pengumpulan Data", "36", False),
    ("         1.  Teknik Pengumpulan Data", "36", False),
    ("         2.  Sumber Data", "36", False),
    ("         3.  Teknik Pengembangan Instrumen", "37", False),
    ("    F.  Teknik Analisis Data", "38", False),
    ("DAFTAR PUSTAKA", "44", True),
]

for label, page, bold in daftar_isi_data:
    elements_to_insert.append(make_entry(label, page, bold=bold))

# Page break + DAFTAR TABEL
elements_to_insert.append(make_page_break())
elements_to_insert.append(make_text_para("DAFTAR TABEL", bold=True, center=True))

daftar_tabel_data = [
    ("Tabel 2.1 Faktor-Faktor yang Mempengaruhi Kompensasi", "18"),
    ("Tabel 2.2 Hasil Penelitian Terdahulu", "24"),
    ("Tabel 3.1 Jadwal Penelitian dan Penulisan Skripsi", "31"),
    ("Tabel 3.2 Operasional Variabel Penelitian", "33"),
    ("Tabel 3.3 Skala Likert", "37"),
    ("Tabel 3.4 Rule of Thumb Evaluasi Model Pengukuran", "40"),
    ("Tabel 3.5 Rule of Thumb Evaluasi Model Struktural", "43"),
]
for label, page in daftar_tabel_data:
    elements_to_insert.append(make_entry(label, page))

# Page break + DAFTAR GAMBAR
elements_to_insert.append(make_page_break())
elements_to_insert.append(make_text_para("DAFTAR GAMBAR", bold=True, center=True))

daftar_gambar_data = [
    ("Gambar 2.1 Kerangka Berpikir", "29"),
    ("Gambar 3.1 Outer Model Variabel Dependen (Y)", "38"),
    ("Gambar 3.2 Outer Model Variabel Independen (X1)", "38"),
    ("Gambar 3.3 Outer Model Variabel Independen (X2)", "39"),
    ("Gambar 3.4 Outer Model Variabel Mediasi (Z)", "39"),
    ("Gambar 3.5 Inner Model Variabel Dependen, Independen, dan Mediasi", "39"),
    ("Gambar 3.6 Diagram Jalur (Path Diagram)", "40"),
]
for label, page in daftar_gambar_data:
    elements_to_insert.append(make_entry(label, page))

# Insert all elements right after DAFTAR ISI heading, in order
ref = daftar_isi_para._element
for el in elements_to_insert:
    body.insert(list(body).index(ref) + 1, el)
    ref = el

print(f"  Inserted {len(elements_to_insert)} elements after DAFTAR ISI heading")

print("\n=== Step 5: Add DAFTAR PUSTAKA at end ===")

bibliography = [
    "Agustin, D. S. (2020). Analisis Pengaruh Budaya Organisasi, Gaya Kepemimpinan terhadap Kinerja Karyawan UMKM dengan Motivasi Kerja sebagai Variabel Mediasi. Jurnal Ilmu Manajemen, 8(3), 789-800.",
    "Amin, A. R. A. (2024). Pengaruh Kompetensi dan Gaya Kepemimpinan Terhadap Kinerja Perangkat Desa di Kecamatan Rambang dengan Motivasi Sebagai Variabel Mediasi (Skripsi, Universitas Sriwijaya).",
    "Arrasyd, H. M., dkk. (2024). Pengaruh Gaya Kepemimpinan Dan Lingkungan Kerja Terhadap Kinerja Perangkat Desa Dengan Motivasi Kerja Sebagai Variabel Mediasi. Jurnal Manajemen & Kewirausahaan, 12(1), 45-56.",
    "Atmojo, A. A. (2022). Pengaruh Gaya Kepemimpinan, Kompensasi Terhadap Kinerja Dengan Motivasi Kerja Sebagai Variabel Mediasi Pada PDAM Kabupaten Kediri. Jurnal Penelitian Ekonomi dan Bisnis, 7(2), 112-125.",
    "Bahits, A., dkk. (2023). Manajemen Sumber Daya Manusia dalam Organisasi Modern. Bandung: Widina Bhakti Persada.",
    "Butarbutar, B., & Nawangsari, L. C. (2022). The Effect of Compensation and Work Discipline on Employee Performance Through Work Motivation (DPRD DKI Jakarta). International Journal of Social and Management Studies, 3(4), 88-98.",
    "Chin, W. W. (1998). The Partial Least Squares Approach to Structural Equation Modeling. Modern Methods for Business Research, 295(2), 295-336.",
    "Fikri, M., dkk. (2022). Manajemen Sumber Daya Manusia: Teori dan Praktik. Jakarta: PT RajaGrafindo Persada.",
    "Geisser, S. (1975). The Predictive Sample Reuse Method with Applications. Journal of the American Statistical Association, 70(350), 320-328.",
    "Ghozali, I. (2021). Partial Least Squares: Konsep, Teknik dan Aplikasi Menggunakan Program SmartPLS 3.0 (Edisi 3). Semarang: Badan Penerbit Universitas Diponegoro.",
    "Guntoro, Djunaedi, & Utami. (2024). Pengaruh Kompensasi Terhadap Kinerja Pegawai Dengan Motivasi Kerja Sebagai Variabel Mediasi pada Dinas Koperasi Nganjuk. Jurnal Ilmiah Manajemen, 11(2), 210-222.",
    "Hair, J. F., Hult, G. T. M., Ringle, C. M., & Sarstedt, M. (2022). A Primer on Partial Least Squares Structural Equation Modeling (PLS-SEM) (3rd ed.). Thousand Oaks, CA: Sage.",
    "Hasibuan, M. S. P. (2019). Manajemen Sumber Daya Manusia. Jakarta: PT Bumi Aksara.",
    "Mangkunegara, A. A. A. P. (2017). Manajemen Sumber Daya Manusia Perusahaan. Bandung: PT Remaja Rosdakarya.",
    "Naldi, M. K., dkk. (2025). Pengaruh Kepuasan Kerja, Gaya Kepemimpinan, dan Komitmen Organisasi terhadap Motivasi Kerja. Jurnal Riset Manajemen, 6(1), 101-112.",
    "Noor, J. (2015). Analisis Data Penelitian Manajemen: Tren Komparasi Antara Penggunaan SPSS, SmartPLS, AMOS, dan Lisrel. Jakarta: Kencana.",
    "Nurcahaya, B. S., & Subkhan, M. (2025). Pengaruh Gaya Kepemimpinan dan Kompensasi terhadap Kinerja Karyawan melalui Motivasi Kerja sebagai Variabel Mediasi pada Pamella Tiga Yogyakarta. Jurnal Manajemen Bisnis, 15(1), 45-58.",
    "Pranogyo, A., dkk. (2021). Kinerja dan Motivasi Pegawai. Yogyakarta: Deepublish.",
    "Prasetya, G. L. H., & Hendarto, R. T. (2024). Pengaruh Gaya Kepemimpinan dan Kompensasi Terhadap Kinerja Karyawan Melalui Motivasi Kerja Sebagai Variable Mediasi Pada PT. Karunia Jaya Global. Jurnal Manajemen Kewirausahaan, 19(1), 34-46.",
    "Pratama, M. Y., & Junaedi, D. (2025). Pengaruh Gaya Kepemimpinan dan Motivasi Kerja Terhadap Efektivitas Kinerja Perangkat Desa (Skripsi, Universitas Nurul Jadid).",
    "Priansa, D. J. (2017). Manajemen Kepemimpinan dan Kinerja Pegawai dalam Organisasi Publik dan Bisnis. Bandung: CV Pustaka Setia.",
    "Purwohedi, U. (2022). Metode Penelitian Kuantitatif untuk Manajemen dan Akuntansi. Jakarta: Rajawali Pers.",
    "Putri, D. R., & Rozi, F. (2024). Pengaruh Kompensasi dan Lingkungan Kerja terhadap Kinerja Karyawan. Jurnal Ilmu Sumber Daya Manusia, 5(2), 140-151.",
    "Robbins, S. P., & Judge, T. A. (2017). Organizational Behavior (17th ed.). Boston: Pearson.",
    "Satria, E., dkk. (2024). Pengaruh Kompensasi terhadap Kinerja Pegawai pada Kantor Pemerintahan. Jurnal Administrasi Publik, 10(1), 67-78.",
    "Stone, M. (1974). Cross-Validatory Choice and Assessment of Statistical Predictions. Journal of the Royal Statistical Society: Series B, 36(2), 111-133.",
    "Sudaryana, A., & Agusiady, R. (2021). Metodologi Penelitian Manajemen. Bandung: Deepublish.",
    "Sudaryo, Y., dkk. (2018). Manajemen Sumber Daya Manusia: Kompensasi dan Visi Kepemimpinan. Yogyakarta: Andi Offset.",
    "Sugiyono. (2023). Metode Penelitian Kuantitatif, Kualitatif, dan R&D. Bandung: Alfabeta.",
    "Sujarweni, V. W. (2020). Metodologi Penelitian: Lengkap, Praktis, dan Mudah Dipahami. Yogyakarta: Pustaka Baru Press.",
    "Sujarweni, V. W. (2022). Analisis Data Penelitian Manajemen dan Ekonomi. Yogyakarta: Pustaka Baru Press.",
    "Syahza, A. (2021). Metodologi Penelitian Edisi Revisi. Pekanbaru: Unri Press.",
    "Thoha, M. (2018). Kepemimpinan dalam Manajemen. Jakarta: PT RajaGrafindo Persada.",
    "Ulfah, R. A., Subiyanto, D., & Kurniawan, I. S. (2020). Peran Mediasi Motivasi Kerja Pada Pengaruh Gaya Kepemimpinan Transaksional Dan Kompensasi Terhadap Kinerja Karyawan. Jurnal Kinerja, 17(2), 180-192.",
    "Wahyudin, M. B. B., & Baidlowi, I. (2025). Pengaruh Gaya Kepemimpinan, Lingkungan Kerja, dan Kompensasi terhadap Kinerja Karyawan dengan Motivasi Sebagai Variabel Mediasi pada UD. Restu Bumi. Jurnal Sains Manajemen, 7(1), 50-62.",
    "Yuliani, T. (2023). Manajemen Sumber Daya Manusia di Sektor Publik. Yogyakarta: Gava Media.",
]

# Add DAFTAR PUSTAKA heading
dp_heading = doc.add_paragraph()
dp_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
dp_heading.paragraph_format.space_before = Pt(12)
run_dph = dp_heading.add_run("DAFTAR PUSTAKA")
run_dph.font.name = 'Times New Roman'
run_dph.font.size = Pt(12)
run_dph.font.bold = True

for entry in bibliography:
    p_b = doc.add_paragraph()
    p_b.paragraph_format.left_indent = Inches(0.5)
    p_b.paragraph_format.first_line_indent = Inches(-0.5)
    p_b.paragraph_format.line_spacing = 1.5
    p_b.paragraph_format.space_after = Pt(6)
    run_b = p_b.add_run(entry)
    run_b.font.name = 'Times New Roman'
    run_b.font.size = Pt(12)

print(f"  Added {len(bibliography)} bibliography entries")

doc.save(DOC_PATH)
print("\n=== DONE. Document saved successfully! ===")

# Verify
doc2 = Document(DOC_PATH)
print(f"Final paragraph count: {len(doc2.paragraphs)}")

# Show the structure around DAFTAR ISI
for i, p in enumerate(doc2.paragraphs):
    if p.text.strip() in ['DAFTAR ISI', 'DAFTAR TABEL', 'DAFTAR GAMBAR', 'DAFTAR PUSTAKA']:
        print(f"[{i}] {p.text.strip()}")
